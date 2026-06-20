#!/usr/bin/env python3
"""Render the Drug Safety revision manuscript as DOCX."""

from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src" / "build_preliminary_manuscript_docx.py"
MANUSCRIPT = ROOT / "manuscript"
PROTOCOL = ROOT / "protocol"
IN = MANUSCRIPT / "drug_safety_revision_manuscript.en.md"
OUT = MANUSCRIPT / "drug_safety_revision_manuscript.docx"
REPORT = PROTOCOL / "drug_safety_revision_docx_report.zh.md"


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def apply_submission_format(path: Path) -> None:
    doc = Document(path)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.different_first_page_header_footer = False
        sect_pr = section._sectPr
        existing = sect_pr.find(qn("w:lnNumType"))
        if existing is not None:
            sect_pr.remove(existing)
        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:restart"), "newPage")
        sect_pr.append(line_numbers)
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.text = "Page "
        add_field(paragraph, "PAGE")
    doc.save(path)


def main() -> None:
    spec = importlib.util.spec_from_file_location("docx_builder", SRC)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    module.IN = IN
    module.OUT = OUT
    module.build_docx()
    apply_submission_format(OUT)
    report = f"""# Drug Safety Word 稿件生成报告

- 已生成：`{OUT.relative_to(ROOT)}`
- 输入：`{IN.relative_to(ROOT)}`
- 已应用双倍行距、页码和按页重启的行号设置，便于投稿前审阅。
- 作者、单位、通讯作者、基金、利益冲突和作者贡献字段已按当前信息补齐。
"""
    REPORT.write_text(report, encoding="utf-8")
    print(report.strip())


if __name__ == "__main__":
    main()
