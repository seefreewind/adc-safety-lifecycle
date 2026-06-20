#!/usr/bin/env python3
"""Render the preliminary manuscript markdown as a Word draft."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
PROTOCOL = ROOT / "protocol"
IN = MANUSCRIPT / "preliminary_manuscript.en.md"
OUT = MANUSCRIPT / "preliminary_manuscript.docx"
REPORT = PROTOCOL / "preliminary_manuscript_docx_report.zh.md"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        raw = lines[idx].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def pretty_table_text(text: str) -> str:
    replacements = {
        "primary_candidate": "primary",
        "sensitivity_candidate": "sensitivity",
        "adverse_event_leading_to_discontinuation": "AE leading to discontinuation",
        "any_adverse_event": "any AE",
        "dose_interruption": "dose interruption",
        "dose_reduction": "dose reduction",
        "fatal_adverse_event": "fatal AE",
        "grade_3_or_higher_adverse_event": "grade >=3 AE",
        "serious_adverse_event": "serious AE",
    }
    return replacements.get(text, text.replace("_", " "))


def add_markdown_table(doc: Document, rows: list[list[str]], landscape: bool = False) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = True
    set_table_width(table, 13680 if landscape else 9360)

    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "F2F4F7")
            text = pretty_table_text(row[c_idx]) if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(7.5 if col_count >= 8 else 9)
            run.font.bold = r_idx == 0
    doc.add_paragraph()


def add_markdown_image(doc: Document, line: str) -> bool:
    match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line.strip())
    if not match:
        return False
    image_path = Path(match.group(1))
    if not image_path.is_absolute():
        image_path = (ROOT / image_path).resolve()
    if not image_path.exists():
        add_paragraph_with_inline_code(doc, f"[Missing image: {image_path}]")
        return True
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    return True


def add_landscape_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def add_portrait_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.strip("`") if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    return paragraph


def build_docx() -> None:
    text = IN.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    style_document(doc)

    idx = 0
    if lines and lines[0].startswith("# "):
        add_title(doc, lines[0][2:].strip())
        idx = 1

    in_tables_section = False
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        if not stripped:
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            wide_table = max(len(row) for row in rows) >= 8 if rows else False
            if wide_table and not in_tables_section:
                add_landscape_section(doc)
                add_markdown_table(doc, rows, landscape=True)
                add_portrait_section(doc)
            else:
                add_markdown_table(doc, rows, landscape=in_tables_section)
            continue
        if stripped.startswith("![") and add_markdown_image(doc, stripped):
            idx += 1
            continue
        if stripped.startswith("## "):
            if stripped == "## Tables" and not in_tables_section:
                add_landscape_section(doc)
                in_tables_section = True
            elif stripped == "## Supplementary Information" and in_tables_section:
                add_portrait_section(doc)
                in_tables_section = False
            doc.add_paragraph(stripped[3:].strip(), style="Heading 1")
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:].strip(), style="Heading 2")
        elif re.match(r"^\d+\.\s", stripped):
            paragraph = add_paragraph_with_inline_code(doc, stripped)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        elif stripped.startswith("- "):
            paragraph = add_paragraph_with_inline_code(doc, stripped[2:], style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
        else:
            add_paragraph_with_inline_code(doc, stripped)
        idx += 1

    doc.core_properties.title = "ADC safety source concordance preliminary manuscript"
    doc.core_properties.subject = "Preliminary manuscript draft"
    doc.save(OUT)


def main() -> None:
    build_docx()
    report = [
        "# Word 初步稿件生成报告",
        "",
        f"- 已生成：`{OUT.relative_to(ROOT)}`",
        "- 样式：standard_business_brief，适合正式学术工作稿。",
        "- 作者栏保持留空。",
        "- 该 DOCX 从 `manuscript/preliminary_manuscript.en.md` 生成。",
    ]
    REPORT.write_text("\n".join(report) + "\n", encoding="utf-8")
    print("\n".join(report))


if __name__ == "__main__":
    main()
